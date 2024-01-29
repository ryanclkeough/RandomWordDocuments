import datetime
import docx
import os
import random
import shutil
import string

def main():
    currentDirectory = os.getcwd()
    folderDateTime = datetime.datetime.now().strftime("%Y_%m_%d-%H_%M_%S")
    newFolderName = f"Documents-{folderDateTime}"
    randomWordDocumentsFolder = os.path.join(currentDirectory, newFolderName)
    try:
        os.mkdir(randomWordDocumentsFolder)
        print(f"Folder '{newFolderName}' created successfully in '{currentDirectory}'!")
    except OSError as error:
        print(f"Error creating folder: {error}")

    randomFileAmount = random.randrange(10, 15)
    for fileNumber in range(randomFileAmount):
        fileName = datetime.datetime.now().strftime("%Y_%m_%d-%H_%M_%S.%f")
        document = docx.Document()
        RandomMathDocument()
        randomAmountOfCharacters = random.randrange(1000, 10000)
        textForDocument = ''.join(random.choices(string.ascii_letters, k=randomAmountOfCharacters))
        document.add_paragraph(textForDocument)
        newDocument = os.path.join(randomWordDocumentsFolder, f"{fileName}.docx")
        document.save(newDocument)

    shutil.make_archive(newFolderName, 'zip', randomWordDocumentsFolder)
    try:
        shutil.rmtree(randomWordDocumentsFolder)
        print(f"Folder '{newFolderName}' and its contents have been deleted successfully.")
    except OSError as error:
        print(f"Error deleting folder: {error}")

def RandomMathDocument():
    mathString = ""
    symbols = ["+", "-", "*", "/", "%", "=", "<", ">", "<=", ">=", " ", "\n"]

    for character in range(18000, 24000):
        randomLetters = random.randrange(1, 3)
        mathString = mathString + random.choices(string.ascii_letters, k=randomLetters)
        mathString = mathString + random.choices(symbols, weights=[1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 5, 1])

    return mathString

if __name__ == '__main__':
  main()